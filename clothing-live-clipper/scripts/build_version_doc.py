# -*- coding: utf-8 -*-
"""Generate Word version changelog for clothing live clip project."""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True)
    return p


def add_para(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def main() -> Path:
    desktop = Path(r"C:\Users\MR\Desktop") / "视频剪辑skill"
    desktop.mkdir(parents=True, exist_ok=True)
    out = desktop / f"版本修改说明_服装带货直播切片_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

    doc = Document()
    # default style
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("服装带货直播智能切片系统")
    set_run_font(r, size=18, bold=True)
    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = title2.add_run("版本修改说明文档")
    set_run_font(r2, size=16, bold=True)

    add_para(doc, f"文档生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
    add_para(doc, "项目名称：clothing-live-clipper + clothing-live-clip Agent Skill")
    add_para(doc, "当前分支：feature/web-video-workstation")
    add_para(doc, "当前最新提交：d858bdc / 09f55aa 等（见版本表）")
    add_para(doc, "仓库：https://github.com/Zzin-cell/clothing-live-clipper")
    add_para(doc, "本机工程：C:\\Users\\MR\\AppData\\grok\\clothing-live-clipper")
    add_para(doc, "Skill 工作区：C:\\Users\\MR\\Desktop\\视频剪辑skill")

    add_heading(doc, "一、产品目标（当前定稿）", 1)
    add_para(doc, "输入：只要直播视频（不要用户上传口播稿）。")
    add_para(doc, "处理：视频 → 自动转写口播并打时间戳 → 只保留衣服相关内容 → 逻辑排序 → 成片。")
    add_para(doc, "输出：约 60 秒切片，1.3 倍速，前 20 秒放最强特点/卖点；不讨论价格；尽量无明显剪辑痕迹。")
    add_para(doc, "执行方式：默认 Agent 直接调用 skill/本地脚本；Web 仅为可选提交台。")

    add_heading(doc, "二、当前能力与路径", 1)
    add_heading(doc, "2.1 默认执行（无前端）", 2)
    add_para(doc, "命令示例：")
    add_para(
        doc,
        r'cd C:\Users\MR\AppData\grok\clothing-live-clipper'
        "\nset PYTHONPATH=src"
        "\nset PATH=%LOCALAPPDATA%\\ffmpeg\\bin;%PATH%"
        "\npython scripts\\agent_clip_video.py \"视频路径.mp4\"",
        size=10,
    )
    add_para(doc, "批量：待剪辑 → 已经完成：python scripts\\batch_desktop_clip.py")

    add_heading(doc, "2.2 关键目录", 2)
    add_para(doc, "本地 Whisper 模型：C:\\Users\\MR\\AppData\\grok\\models\\whisper-tiny")
    add_para(doc, "Agent 产出：clothing-live-clipper\\output\\agent_jobs\\")
    add_para(doc, "桌面成品示例：C:\\Users\\MR\\Desktop\\检查文件\\已经完成\\")
    add_para(doc, "Skill 包：C:\\Users\\MR\\Desktop\\视频剪辑skill\\clothing-live-clip\\")
    add_para(doc, "规范书：C:\\Users\\MR\\Desktop\\视频剪辑skill\\服装带货直播切片Skill规范书.md")

    add_heading(doc, "三、版本修改记录（按时间）", 1)
    add_para(doc, "说明：时间为提交时间；内容为该版本核心变更。")

    versions = [
        ("2026-07-18 18:23", "6c5094a", "v0.1", "Skill 设计规格文档落地（服装直播切片规则初稿）。"),
        ("2026-07-18 18:29", "9aeebc1", "v0.1.1", "Skill 实现计划文档落地。"),
        ("2026-07-18 23:59", "aaec82f", "v0.2", "扩展直播 CTA 词表（小黄车/号链接等，先作 price 类）。"),
        ("2026-07-19 00:03", "826bba1", "v0.3", "同步完整 clipper 工程与 skill 到仓库。"),
        ("2026-07-19 00:12~00:14", "07c0002 / 095875f", "v0.4", "Web 工作台设计与实现计划。"),
        ("2026-07-19 00:16~00:22", "0addc2f~b3d4386", "v0.5", "Web API/UI 视频优先、状态页、README。"),
        ("2026-07-19 00:27~01:08", "877b4e3~998d920", "v0.5.1", "Web UI 白/Apple 风格优化。"),
        ("2026-07-19 01:02", "df662e3", "v0.5.2", "ffmpeg 安装辅助与 start-web 脚本。"),
        ("2026-07-19 01:18", "9f9230a", "v0.6", "Web 侧 Whisper API 自动听写（后因中转无 whisper 模型废弃为主路径）。"),
        ("2026-07-19 01:23~01:32", "eed8dee / 5345fe9", "v0.6.1", "系统状态栏与 API 设置抽屉。"),
        ("2026-07-19 01:39", "7c18e6a", "v0.6.2", "固定加载项目 .env。"),
        ("2026-07-19 02:24", "c601933", "v0.7", "架构调整：Web 只入队，Agent+Skill 处理队列。"),
        ("2026-07-20 15:41", "ea32016", "v0.8", "Web 输入强制只要视频（去掉口播稿/示例入口）。"),
        ("2026-07-20 15:49~16:09", "cf7787f / 3a6bfd2", "v0.9", "Agent 本地 worker：视频→ASR→切片；默认无 Web/无云端 API。"),
        ("2026-07-20 16:36", "8cbc793", "v0.9.1", "本地下载并接入 faster-whisper tiny 模型。"),
        ("2026-07-20 16:54", "3ce3500", "v0.10", "学习复盘：只留衣服相关，剔除“过一下”等直播控场词与零食闲聊。"),
        ("2026-07-20 20:33", "989b3c6", "v0.11", "成片目标时长 55–60s 填充策略。"),
        ("2026-07-20 20:38~20:53", "4d46b12 / 8406eb8", "v0.12", "拼接顺滑尝试；1.3 倍速且终片约 60s。"),
        ("2026-07-20 22:07", "ce714a7", "v0.13", "切片排序逻辑化，逻辑优先于去重。"),
        ("2026-07-20 22:22", "9f8a979", "v0.14", "硬规则：成片不要讨论价格（价格/优惠/挂车话术剔除）。"),
        ("2026-07-20 22:48~22:55", "5624853 / 09f55aa", "v0.15", "先试剪映风转场，后改“直接接”、无可见剪辑痕迹。"),
        ("2026-07-20 23:07", "d858bdc", "v0.16", "前 20 秒强化特点/卖点；加强时间链连贯性。"),
        (datetime.now().strftime("%Y-%m-%d %H:%M"), "（本整理）", "v0.16.1", "完整同步整理：Skill/规范书/本 Word 版本说明；统一当前规则与路径。"),
    ]

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["时间", "版本/提交", "版本号", "修改内容"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True)

    for t, commit, ver, content in versions:
        row = table.add_row().cells
        row[0].text = t
        row[1].text = commit
        row[2].text = ver
        row[3].text = content
        for cell in row:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)

    add_heading(doc, "四、当前生效规则清单（v0.16）", 1)
    rules = [
        "输入只要视频；口播稿由本地 ASR 自动生成并带时间戳。",
        "只保留衣服相关内容；剔除直播控场词（如“过一下”）与零食等无关话题。",
        "不讨论价格（价格/优惠/挂车相关一律不进成片）。",
        "排序逻辑优先：卖点/特点 → 版型 → 面料 → 细节 → 搭配；逻辑大于硬去重。",
        "前约 20 秒强制优先特点/卖点（feature-first）。",
        "最终成片约 55–65 秒（目标 60 秒），默认 1.3 倍速。",
        "拼接策略：直接接（无花式转场），仅极短消爆音，减少剪辑痕迹。",
        "默认 Agent 本地执行；Web 为可选提交入口。",
    ]
    for i, rule in enumerate(rules, 1):
        add_para(doc, f"{i}. {rule}")

    add_heading(doc, "五、同步与部署说明", 1)
    add_para(doc, "1）Skill 已同步到：")
    add_para(doc, r"   - C:\Users\MR\Desktop\视频剪辑skill\clothing-live-clip")
    add_para(doc, r"   - C:\Users\MR\.agents\skills\clothing-live-clip")
    add_para(doc, r"   - C:\Users\MR\AppData\grok\skills\clothing-live-clip")
    add_para(doc, "2）规范书：视频剪辑skill\\服装带货直播切片Skill规范书.md")
    add_para(doc, "3）代码主仓分支 feature/web-video-workstation（本地可能超前 origin，网络恢复后 push）。")
    add_para(doc, "4）测试：pytest 相关用例已通过（约 22 passed）。")

    add_heading(doc, "六、已知限制", 1)
    add_para(doc, "1）本地 ASR 使用 tiny 模型，个别识别不准，可能影响选句质量。")
    add_para(doc, "2）中转 API 无 whisper 通道，不能依赖云端听写。")
    add_para(doc, "3）素材本身服装有效口播不足时，可能略短于 60 秒。")
    add_para(doc, "4）“无明显剪辑痕迹”依赖口播时间链；语义跳跃无法单靠转场消除。")

    add_heading(doc, "七、建议后续版本", 1)
    add_para(doc, "v0.17：升级 ASR 到 base/small，提高口播准确率。")
    add_para(doc, "v0.18：切点对齐句尾/停顿，进一步降低跳切感。")
    add_para(doc, "v0.19：多产品自动分段。")
    add_para(doc, "v0.20：可选导出剪映草稿。")

    add_heading(doc, "八、签收信息", 1)
    add_para(doc, "整理人：Agent（Cursor/Grok 协作实现）")
    add_para(doc, f"整理完成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    add_para(doc, "文档用途：版本归档、交接、回溯修改内容与时间。")

    doc.save(out)
    # also stable name
    stable = desktop / "版本修改说明_服装带货直播切片_最新.docx"
    doc.save(stable)
    print(out)
    print(stable)
    return out


if __name__ == "__main__":
    main()
